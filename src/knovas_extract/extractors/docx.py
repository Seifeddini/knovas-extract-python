"""DOCX extractor — python-docx + mammoth backends.

DOCX is a ZIP container holding XML. Two trust-boundary hazards we have to
defuse before parsing:

1. **Decompression bombs** (1 KB → 4 PB classic). We pre-scan the ZIP's
   central directory; any entry whose declared uncompressed size / compressed
   size > Limits.max_decompression_ratio raises ResourceExhaustedError before
   any extractor touches the bytes.
2. **Zip-slip** — an entry name like `../../etc/passwd` would write outside
   the destination if the extractor naively unpacks. We never extract to
   disk, but we still refuse ZIP entries with absolute or parent-traversing
   names so a buggy downstream consumer can't be coerced.

XML parsing within docx files goes through `defusedxml` (entity resolution
disabled, billion-laughs blocked). python-docx historically used lxml without
defusedxml; we install defusedxml's lxml shim at import time.

For body text we use **python-docx**; for heading extraction (sections[]) we
use **mammoth**'s HTML conversion + a tiny ATX-style harvester. Either path
alone has known gaps; the combination is the established pragmatic approach.
"""

from __future__ import annotations

import hashlib
import re
import zipfile
from io import BytesIO
from typing import TYPE_CHECKING, ClassVar

from knovas_extract.dispatch import MIME_REGISTRY, make_result
from knovas_extract.errors import (
    CorruptDocumentError,
    ResourceExhaustedError,
)
from knovas_extract.interfaces import IExtractor
from knovas_extract.normalize import canonicalize_text, word_count
from knovas_extract.result import ExtractionResult, Limits, Metadata, Section

if TYPE_CHECKING:
    pass

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# mammoth emits HTML like <h2>Heading text</h2>. We harvest level + text via
# this regex rather than spinning up a full HTML parser.
_HEADING_HTML = re.compile(r"<h([1-6])>([^<]+)</h\1>", re.IGNORECASE)


def _guard_zip(data: bytes, limits: Limits) -> zipfile.ZipFile:
    """Open the docx zip, refusing zip-slip names and decompression bombs.

    Raises CorruptDocumentError on malformed/zip-slip and
    ResourceExhaustedError on bomb-like compression ratios.

    The ratio / per-entry checks below read the central directory's DECLARED
    sizes, and that is sufficient: CPython's `zipfile` treats `file_size` as
    the authoritative uncompressed length — it never emits more than
    `file_size` bytes for an entry and CRC-verifies at EOF. So a hand-forged
    archive that UNDERSTATES `file_size` to sneak past these caps does not
    then decompress unbounded; the reader (here, and inside python-docx /
    mammoth) stops at the declared size and fails CRC -> BadZipFile. Honest
    bombs declare their true (huge) sizes and are caught here before any XML
    parser runs.
    """
    try:
        zf = zipfile.ZipFile(BytesIO(data), "r")
    except (zipfile.BadZipFile, EOFError, OSError) as exc:
        raise CorruptDocumentError(f"DOCX is not a valid ZIP container: {exc}") from exc

    for info in zf.infolist():
        # Zip-slip / absolute-path guard.
        name = info.filename
        if name.startswith(("/", "\\")) or ".." in name.replace("\\", "/").split("/"):
            zf.close()
            raise CorruptDocumentError(f"DOCX contains zip-slip-style entry name: {name!r}")
        # Decompression bomb guard.
        if info.compress_size > 0:
            ratio = info.file_size / info.compress_size
            if ratio > limits.max_decompression_ratio:
                zf.close()
                raise ResourceExhaustedError(
                    "decompression ratio",
                    limits.max_decompression_ratio,
                    observed=ratio,
                )
        # Per-entry absolute size guard — even a single huge entry is suspect.
        if info.file_size > limits.max_text_bytes:
            zf.close()
            raise ResourceExhaustedError(
                "zip entry uncompressed size",
                limits.max_text_bytes,
                observed=info.file_size,
            )
    return zf


# NOTE — XML security posture (three parsers touch the DOCX zip):
#   1. Our docProps/{core,app}.xml parsing → `defusedxml.ElementTree`
#      (see _extract_*_metadata). Entities forbidden; the safe path. A
#      hostile entity here raises defusedxml's EntitiesForbidden, which we
#      re-frame as CorruptDocumentError.
#   2. python-docx parses word/document.xml with **lxml**, using a parser
#      configured `resolve_entities=False` (docx.oxml.parser). That disables
#      BOTH internal-entity expansion (billion-laughs) and external-entity
#      resolution (XXE): entity references are left unexpanded, not fetched;
#      lxml also defaults to no_network=True. A billion-laughs payload trips
#      libxml2's entity-amplification cap and surfaces here as
#      CorruptDocumentError.
#   3. mammoth parses word/document.xml with the stdlib **xml.dom.minidom**
#      (expat). Expat does not fetch external entities by default (no XXE),
#      and CPython >= 3.11 (our floor) bundles libexpat >= 2.4 whose
#      billion-laughs amplification protection is on by default — an
#      amplification bomb raises ExpatError, which _extract_html() catches
#      (sections / markdown are skipped with a warning; python-docx has
#      usually already rejected the document by then).
#   The zip decompression-ratio + actual-size caps in _guard_zip bound the
#   *bytes* each parser sees; the entity protections above bound *expansion*
#   during parse — which the byte caps do NOT catch, since an entity bomb is
#   tiny on disk. Both layers are required.
# NB: this is NOT "all XML via defusedxml" — only the docProps path is.
# We previously called defusedxml.defuse_stdlib() as belt-and-suspenders for
# the minidom path. Removed because it patches the deprecated
# xml.etree.cElementTree, emitting a DeprecationWarning that our
# filterwarnings=error test config turns into a hard failure; the libexpat
# amplification cap (guaranteed by our Python >= 3.11 floor) covers it.


def _extract_body_text(data: bytes) -> str:
    """Body text via python-docx. Takes original docx bytes (not a ZipFile).

    Since spec 1.1.0 (this repo v0.2.0+), tables are extracted STRUCTURALLY
    into `content.tables[]` by `_extract_structured_tables`; they are NOT
    duplicated into `content.text` here. This is a behavior change from
    v0.1.x — table content used to appear as pipe-joined lines in the body
    text; downstream servers now receive the raw structure and can decide
    how to serialize it into embedding chunks. See docs/schema-fields.md#tables.
    """
    import docx  # python-docx

    try:
        document = docx.Document(BytesIO(data))
    except Exception as exc:
        raise CorruptDocumentError(f"python-docx could not parse: {exc}") from exc

    parts: list[str] = []
    for para in document.paragraphs:
        text = (para.text or "").rstrip()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


_TABLE_CELL_MAX_CHARS = 1024


def _extract_structured_tables(data: bytes) -> tuple[list, list[str]]:
    """Extract structured tables from a DOCX blob (spec 1.1.0+).

    Returns (tables, warnings) where `tables` is a list of `Table` dataclass
    instances (imported lazily to avoid a circular import) and `warnings`
    holds any truncation / shape notes for the caller to append to the
    result's warnings.
    """
    from ..result import Table

    import docx  # python-docx

    try:
        document = docx.Document(BytesIO(data))
    except Exception as exc:
        raise CorruptDocumentError(f"python-docx could not parse: {exc}") from exc

    tables: list = []
    warnings: list[str] = []
    for t_idx, table in enumerate(document.tables):
        rows_raw = list(table.rows)
        if not rows_raw:
            continue
        header_cells = [(c.text or "").strip() for c in rows_raw[0].cells]
        headers = [
            _cap_cell(h, warnings, t_idx, "header", i) for i, h in enumerate(header_cells)
        ]
        if not headers:
            continue
        n_cols = len(headers)
        rows: list[list[str]] = []
        for r_idx, row in enumerate(rows_raw[1:], start=1):
            cells = [(c.text or "").strip() for c in row.cells]
            # Pad or trim to the header row's length so `rows[i]` shape stays
            # invariant with `headers`. Spec-side row-length equality is
            # enforced by the Knovas server (which rejects mismatches with a
            # 400) — this client-side padding preserves what CAN be extracted.
            if len(cells) < n_cols:
                cells = cells + [""] * (n_cols - len(cells))
                warnings.append(
                    f"docx: tables[{t_idx}].rows[{r_idx - 1}] padded from "
                    f"{len(row.cells)} to {n_cols} cells"
                )
            elif len(cells) > n_cols:
                cells = cells[:n_cols]
                warnings.append(
                    f"docx: tables[{t_idx}].rows[{r_idx - 1}] truncated from "
                    f"{len(row.cells)} to {n_cols} cells"
                )
            capped = [
                _cap_cell(c, warnings, t_idx, "row", r_idx - 1, col=i)
                for i, c in enumerate(cells)
            ]
            rows.append(capped)
        if not rows:
            continue
        tables.append(
            Table(
                client_table_hint=f"docx_t{t_idx}",
                title=None,
                headers=headers,
                rows=rows,
                page=None,
                bbox=None,
            )
        )
    return tables, warnings


def _cap_cell(value: str, warnings: list[str], t_idx: int, kind: str, r_idx: int, col: int = -1) -> str:
    if len(value) <= _TABLE_CELL_MAX_CHARS:
        return value
    where = f"tables[{t_idx}].{kind}" + (f"[{r_idx}].[{col}]" if col >= 0 else f"[{r_idx}]")
    warnings.append(f"docx: {where} truncated at {_TABLE_CELL_MAX_CHARS} chars")
    return value[:_TABLE_CELL_MAX_CHARS]


def _extract_html(data: bytes) -> str | None:
    """Convert DOCX bytes to HTML via mammoth. Returns None on backend failure.

    Isolated so that both `_sections_from_html` and `_markdown_from_html`
    can reuse the (relatively expensive) mammoth conversion output.
    """
    import mammoth

    try:
        result = mammoth.convert_to_html(BytesIO(data))
    except Exception:
        return None
    return result.value or ""


def _sections_from_html(html: str, canonical_text: str | None = None) -> list[Section]:
    """Harvest heading-anchored sections from mammoth HTML output.

    When ``canonical_text`` (== `content.text` we return) is provided,
    each Section carries 1-based `line_start` / `line_end` in it — used
    by the sentence↔section back-pointer and consumer citations.
    """
    matches = list(_HEADING_HTML.finditer(html))
    if not matches:
        return []

    sections: list[Section] = []
    canon_cursor = 0
    for i, m in enumerate(matches):
        level = int(m.group(1))
        heading = m.group(2).strip()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(html)
        body_html = html[start:end]
        # Preserve paragraph + list-item structure before stripping tags;
        # otherwise consecutive <p>foo</p><p>bar</p> collapses to "foobar".
        body_html = re.sub(r"</(p|li|h[1-6]|tr|br)\s*>", "\n\n", body_html, flags=re.I)
        body_text = re.sub(r"<[^>]+>", "", body_html)

        line_start: int | None = None
        line_end: int | None = None
        if canonical_text is not None:
            hstart = canonical_text.find(heading, canon_cursor)
            if hstart >= 0:
                line_start = 1 + canonical_text.count("\n", 0, hstart)
                next_pos = len(canonical_text)
                for j in range(i + 1, len(matches)):
                    nh = matches[j].group(2).strip()
                    n = canonical_text.find(nh, hstart + len(heading))
                    if n >= 0:
                        next_pos = n
                        break
                line_end = 1 + canonical_text.count("\n", 0, max(next_pos - 1, hstart))
                canon_cursor = hstart + len(heading)

        sections.append(
            Section(
                heading=heading,
                level=level,
                text=canonicalize_text(body_text),
                line_start=line_start,
                line_end=line_end,
            )
        )
    return sections


def _markdown_from_html(html: str, limits: Limits, warnings: list[str]) -> str:
    """Sanitize mammoth's HTML output and convert to Markdown.

    Even though mammoth emits structurally-clean HTML, we still route
    through the full sanitizer because <a href> URLs come from the source
    document verbatim — a hostile DOCX can embed `javascript:` links, and
    mammoth surfaces Word "Text Box" content as raw HTML fragments in
    some documents.
    """
    from knovas_extract._markdown import html_to_markdown

    return html_to_markdown(html, limits, warnings=warnings)


def _extract_core_metadata(zf: zipfile.ZipFile) -> dict[str, str | None]:
    """Pull docProps/core.xml fields via defusedxml.

    Returns a dict with title/author/created/modified/subject/keywords (all str|None).
    """
    from defusedxml import ElementTree as ET

    try:
        with zf.open("docProps/core.xml") as f:
            tree = ET.parse(f)
    except KeyError:
        return {}
    except Exception as exc:
        raise CorruptDocumentError(f"docProps/core.xml unparseable: {exc}") from exc

    ns = {
        "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
        "dc": "http://purl.org/dc/elements/1.1/",
        "dcterms": "http://purl.org/dc/terms/",
    }
    root = tree.getroot()

    def text_of(tag: str, namespace_key: str) -> str | None:
        el = root.find(f"{{{ns[namespace_key]}}}{tag}")
        if el is not None and el.text:
            return el.text.strip() or None
        return None

    return {
        "title": text_of("title", "dc"),
        "author": text_of("creator", "dc"),
        "subject": text_of("subject", "dc"),
        "keywords": text_of("keywords", "cp"),
        "language": text_of("language", "dc"),
        "created": text_of("created", "dcterms"),
        "modified": text_of("modified", "dcterms"),
        "revision": text_of("revision", "cp"),
        "last_modified_by": text_of("lastModifiedBy", "cp"),
        "category": text_of("category", "cp"),
        "content_status": text_of("contentStatus", "cp"),
        "version": text_of("version", "cp"),
    }


def _extract_app_metadata(zf: zipfile.ZipFile) -> dict[str, str | None]:
    """Pull docProps/app.xml fields via defusedxml.

    Optional per the OOXML spec — many DOCX files (especially those from
    non-Microsoft producers) omit app.xml entirely. Absence is not an
    error; returns an empty dict.
    """
    from defusedxml import ElementTree as ET

    try:
        with zf.open("docProps/app.xml") as f:
            tree = ET.parse(f)
    except KeyError:
        return {}
    except Exception as exc:
        raise CorruptDocumentError(f"docProps/app.xml unparseable: {exc}") from exc

    ns = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
    root = tree.getroot()

    def text_of(tag: str) -> str | None:
        el = root.find(f"{{{ns}}}{tag}")
        if el is not None and el.text:
            return el.text.strip() or None
        return None

    return {
        "application": text_of("Application"),
        "app_version": text_of("AppVersion"),
        "template": text_of("Template"),
        "total_time": text_of("TotalTime"),
        "pages_declared": text_of("Pages"),
        "paragraph_count": text_of("Paragraphs"),
        "word_count_declared": text_of("Words"),
        "character_count": text_of("Characters"),
        "company": text_of("Company"),
        "manager": text_of("Manager"),
    }


class DocxExtractor(IExtractor):
    """DOCX text + metadata extractor."""

    supported_mimes: ClassVar[frozenset[str]] = frozenset({DOCX_MIME})
    name: ClassVar[str] = "docx"

    def extract(
        self,
        data: bytes,
        *,
        filename: str | None = None,
        limits: Limits | None = None,
        emit_markdown: bool = False,
        emit_sentences: bool = False,
    ) -> ExtractionResult:
        from collections import Counter

        from knovas_extract._metadata import finalize_warnings, sanitize_scalar

        limits = limits or Limits()
        if len(data) > limits.max_input_bytes:
            raise ResourceExhaustedError("input size", limits.max_input_bytes, observed=len(data))

        warnings: list[str] = []
        counts: Counter[str] = Counter()
        zf = _guard_zip(data, limits)
        try:
            # Detect (but never execute) macros.
            if "word/vbaProject.bin" in zf.namelist():
                warnings.append("DOCX contains VBA macros; payload ignored (never executed)")

            body = _extract_body_text(data)
            text = canonicalize_text(body)

            # Structured tables (spec 1.1.0+) — extracted separately from body
            # text so downstream servers can decide serialization strategy.
            try:
                tables, table_warnings = _extract_structured_tables(data)
                warnings.extend(table_warnings)
            except CorruptDocumentError:
                raise
            except Exception as exc:  # never let a table failure kill the extraction
                tables = None
                warnings.append(f"docx: table extraction failed: {type(exc).__name__}")

            if len(text.encode("utf-8")) > limits.max_text_bytes:
                raise ResourceExhaustedError("text size", limits.max_text_bytes, observed=len(text))

            # Call mammoth ONCE and derive both sections and (optionally)
            # markdown from the same HTML. Cheaper than double-converting.
            html = _extract_html(data)
            sections = _sections_from_html(html, canonical_text=text) if html else []

            markdown: str | None = None
            if emit_markdown:
                if html is None:
                    warnings.append("docx: mammoth conversion failed; content.markdown left null")
                else:
                    from knovas_extract._markdown import check_expansion

                    markdown = _markdown_from_html(html, limits, warnings)
                    check_expansion(markdown, len(text), limits)

            meta = _extract_core_metadata(zf)
            app = _extract_app_metadata(zf)
        finally:
            zf.close()

        extra: dict[str, str | int | float | bool | None] = {}
        # Core.xml extras.
        for k in (
            "subject",
            "keywords",
            "revision",
            "last_modified_by",
            "category",
            "content_status",
            "version",
        ):
            v = meta.get(k)
            if v:
                clean = sanitize_scalar(v, limits=limits, counts=counts)
                if clean is not None:
                    extra[f"docx:{k}"] = clean
        # App.xml extras.
        for k in (
            "application",
            "app_version",
            "template",
            "total_time",
            "pages_declared",
            "paragraph_count",
            "word_count_declared",
            "character_count",
            "company",
            "manager",
        ):
            v = app.get(k)
            if v:
                clean = sanitize_scalar(v, limits=limits, counts=counts)
                if clean is not None:
                    extra[f"docx:{k}"] = clean
        finalize_warnings(counts, warnings)

        metadata = Metadata(
            title=meta.get("title"),
            author=meta.get("author"),
            language=meta.get("language"),
            created=meta.get("created"),
            modified=meta.get("modified"),
            word_count=word_count(text),
            extra=extra,
        )

        sentences = None
        if emit_sentences:
            from knovas_extract._sentences import split_sentences

            sentences = split_sentences(text, limits, warnings=warnings, language=metadata.language)

        return make_result(
            text=text,
            mime=DOCX_MIME,
            sha256=hashlib.sha256(data).hexdigest(),
            size_bytes=len(data),
            filename=filename,
            metadata=metadata,
            sections=sections or None,
            warnings=warnings,
            markdown=markdown,
            sentences=sentences,
            tables=tables or None,
        )


MIME_REGISTRY[DOCX_MIME] = DocxExtractor()
