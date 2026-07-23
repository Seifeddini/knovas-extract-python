"""Unit tests for structured-table extraction from DOCX (spec 1.1.0+).

Verifies the behavior change from v0.1.x:
  * Tables now populate `content.tables[]` instead of being pipe-joined
    into `content.text`.
  * Row length is normalized to match `headers` length (pad or trim + warn).
  * `client_table_hint` follows the `docx_t{index}` convention.
  * Cells > 1024 chars are truncated with a warning.
"""

from __future__ import annotations

import io

import pytest

from knovas_extract import extract

pytest.importorskip("docx")
import docx as _docx  # python-docx  # noqa: E402


def _make_docx_with_table(headers: list[str], rows: list[list[str]]) -> bytes:
    doc = _docx.Document()
    doc.add_paragraph("Some prose before the table.")
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    for c, h in enumerate(headers):
        t.rows[0].cells[c].text = h
    for r, row in enumerate(rows):
        for c, v in enumerate(row):
            t.rows[r + 1].cells[c].text = v
    doc.add_paragraph("Some prose after the table.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_no_tables() -> bytes:
    doc = _docx.Document()
    doc.add_paragraph("Only prose.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# -------------------- structural extraction --------------------


@pytest.mark.unit
def test_docx_table_appears_in_content_tables() -> None:
    data = _make_docx_with_table(
        headers=["Vendor", "Invoice", "Amount"],
        rows=[
            ["Acme", "12345", "500.00"],
            ["Bosch", "12346", "1200.00"],
        ],
    )
    r = extract(
        data, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert r.content.tables is not None
    assert len(r.content.tables) == 1
    t = r.content.tables[0]
    assert t.client_table_hint == "docx_t0"
    assert t.headers == ["Vendor", "Invoice", "Amount"]
    assert t.rows == [
        ["Acme", "12345", "500.00"],
        ["Bosch", "12346", "1200.00"],
    ]
    assert t.page is None
    assert t.bbox is None


@pytest.mark.unit
def test_docx_table_no_longer_appears_in_body_text() -> None:
    """Behavior change from v0.1.x — tables are structural, not prose."""
    data = _make_docx_with_table(
        headers=["Vendor", "Invoice", "Amount"],
        rows=[
            ["UniqueVendorNameForAssertion", "12345", "500.00"],
        ],
    )
    r = extract(
        data, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    # Body prose must still be present.
    assert "prose before the table" in r.content.text
    assert "prose after the table" in r.content.text
    # But table cell content must NOT appear in content.text anymore.
    assert "UniqueVendorNameForAssertion" not in r.content.text
    assert "12345" not in r.content.text


@pytest.mark.unit
def test_docx_no_tables_emits_none() -> None:
    r = extract(
        _make_docx_no_tables(),
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert r.content.tables is None


@pytest.mark.unit
def test_docx_multiple_tables_get_indexed_hints() -> None:
    doc = _docx.Document()
    for _ in range(3):
        t = doc.add_table(rows=2, cols=2)
        t.rows[0].cells[0].text = "H1"
        t.rows[0].cells[1].text = "H2"
        t.rows[1].cells[0].text = "v1"
        t.rows[1].cells[1].text = "v2"
    buf = io.BytesIO()
    doc.save(buf)
    r = extract(
        buf.getvalue(),
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    hints = [t.client_table_hint for t in (r.content.tables or [])]
    assert hints == ["docx_t0", "docx_t1", "docx_t2"]


# -------------------- bounds & sanitation --------------------


@pytest.mark.unit
def test_docx_oversized_cell_truncated_with_warning() -> None:
    long_value = "x" * 2000
    data = _make_docx_with_table(headers=["h"], rows=[[long_value]])
    r = extract(
        data,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert r.content.tables is not None
    t = r.content.tables[0]
    assert len(t.rows[0][0]) == 1024
    assert any("truncated" in w for w in r.warnings)


@pytest.mark.unit
def test_docx_serializes_to_dict_with_tables() -> None:
    """to_dict must include content.tables[] on the JSON round-trip."""
    data = _make_docx_with_table(headers=["a", "b"], rows=[["1", "2"]])
    r = extract(
        data,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    d = r.to_dict()
    assert d["content"]["tables"] is not None
    t = d["content"]["tables"][0]
    assert t["client_table_hint"] == "docx_t0"
    assert t["headers"] == ["a", "b"]
    assert t["rows"] == [["1", "2"]]
    assert t["title"] is None
    assert t["page"] is None
    assert t["bbox"] is None
