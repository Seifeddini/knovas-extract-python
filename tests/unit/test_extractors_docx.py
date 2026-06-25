"""Unit tests for the DOCX extractor — edge cases not covered by the golden corpus."""

from __future__ import annotations

import io
import zipfile

import pytest

from knovas_extract import extract
from knovas_extract.errors import (
    CorruptDocumentError,
    ResourceExhaustedError,
)
from knovas_extract.result import Limits

pytest.importorskip("docx")  # extras [docx] not installed → skip module
import docx  # noqa: E402

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@pytest.fixture
def simple_docx_bytes() -> bytes:
    d = docx.Document()
    d.add_paragraph("Hello from docx.")
    d.add_paragraph("Second paragraph.")
    d.core_properties.title = "Unit Doc"
    d.core_properties.author = "Unit Tester"
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


@pytest.fixture
def headed_docx_bytes() -> bytes:
    d = docx.Document()
    d.add_heading("Top", level=1)
    d.add_paragraph("Body 1.")
    d.add_heading("Sub", level=2)
    d.add_paragraph("Body 2.")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


@pytest.mark.unit
def test_extracts_text_and_metadata(simple_docx_bytes: bytes) -> None:
    r = extract(simple_docx_bytes, mime=DOCX_MIME)
    assert "Hello from docx." in r.content.text
    assert "Second paragraph." in r.content.text
    assert r.metadata.title == "Unit Doc"
    assert r.metadata.author == "Unit Tester"


@pytest.mark.unit
def test_extracts_sections(headed_docx_bytes: bytes) -> None:
    r = extract(headed_docx_bytes, mime=DOCX_MIME)
    assert r.content.sections is not None
    by_heading = {s.heading: s for s in r.content.sections}
    assert by_heading["Top"].level == 1
    assert by_heading["Sub"].level == 2
    assert "Body 1." in by_heading["Top"].text
    assert "Body 2." in by_heading["Sub"].text


@pytest.mark.unit
def test_not_a_zip_raises_corrupt() -> None:
    with pytest.raises(CorruptDocumentError):
        extract(b"\x00 not a docx at all", mime=DOCX_MIME)


@pytest.mark.unit
def test_zip_with_zero_relationships_raises_corrupt() -> None:
    """An empty valid ZIP isn't a valid DOCX — python-docx fails."""
    buf = io.BytesIO()
    z = zipfile.ZipFile(buf, "w")
    z.writestr("dummy.txt", b"hello")
    z.close()
    with pytest.raises(CorruptDocumentError):
        extract(buf.getvalue(), mime=DOCX_MIME)


@pytest.mark.unit
def test_zip_slip_path_rejected() -> None:
    buf = io.BytesIO()
    z = zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED)
    z.writestr("../../etc/passwd-fake", b"x")
    z.close()
    with pytest.raises(CorruptDocumentError, match="zip-slip"):
        extract(buf.getvalue(), mime=DOCX_MIME)


@pytest.mark.unit
def test_absolute_path_rejected() -> None:
    buf = io.BytesIO()
    z = zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED)
    # Build a ZipInfo with an absolute leading slash; some ZIP creators do this.
    info = zipfile.ZipInfo("/abs/path-fake")
    info.compress_type = zipfile.ZIP_DEFLATED
    z.writestr(info, b"x")
    z.close()
    with pytest.raises(CorruptDocumentError, match="zip-slip"):
        extract(buf.getvalue(), mime=DOCX_MIME)


@pytest.mark.unit
def test_decompression_ratio_cap_enforced() -> None:
    payload = b"A" * (1 << 20)  # 1 MiB of one byte → ~1 KiB compressed
    buf = io.BytesIO()
    z = zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9)
    info = zipfile.ZipInfo("bomb.bin")
    info.compress_type = zipfile.ZIP_DEFLATED
    z.writestr(info, payload)
    z.close()
    with pytest.raises(ResourceExhaustedError) as exc:
        extract(buf.getvalue(), mime=DOCX_MIME, limits=Limits(max_decompression_ratio=50))
    assert exc.value.what == "decompression ratio"


@pytest.mark.unit
def test_per_entry_size_cap_enforced() -> None:
    """Even with a 'safe' ratio, a single huge uncompressed entry is rejected."""
    big = b"X" * (2 << 20)  # 2 MiB
    buf = io.BytesIO()
    z = zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_STORED)  # no compression → ratio ~1
    z.writestr("big.bin", big)
    z.close()
    with pytest.raises(ResourceExhaustedError) as exc:
        extract(buf.getvalue(), mime=DOCX_MIME, limits=Limits(max_text_bytes=1 << 20))
    assert exc.value.what == "zip entry uncompressed size"


@pytest.mark.unit
def test_dispatch_routes_docx_when_libmagic_says_zip(simple_docx_bytes) -> None:
    """libmagic detects DOCX as application/zip; dispatch must override via extension."""
    import tempfile
    from pathlib import Path

    # Use the path-based form so the filename extension is available.
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        f.write(simple_docx_bytes)
        path = Path(f.name)
    try:
        r = extract(path)  # no mime override
        assert r.source.mime_type == DOCX_MIME
        assert "Hello from docx." in r.content.text
    finally:
        path.unlink()
