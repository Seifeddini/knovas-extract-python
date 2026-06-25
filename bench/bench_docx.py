"""pytest-benchmark suite for the DOCX extractor."""

from __future__ import annotations

import io

import pytest

pytest.importorskip("docx")
import docx  # noqa: E402

from knovas_extract import extract  # noqa: E402

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _make_docx(paragraphs: int) -> bytes:
    d = docx.Document()
    for i in range(paragraphs):
        if i % 10 == 0:
            d.add_heading(f"Section {i // 10 + 1}", level=2)
        d.add_paragraph(f"Paragraph {i + 1}: " + ("Lorem ipsum dolor sit amet. " * 15))
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


SMALL = _make_docx(5)
MEDIUM = _make_docx(50)
LARGE = _make_docx(500)


@pytest.mark.parametrize(
    "payload",
    [
        pytest.param(SMALL, id="docx-5para"),
        pytest.param(MEDIUM, id="docx-50para"),
        pytest.param(LARGE, id="docx-500para"),
    ],
)
def test_extract_docx(benchmark, payload: bytes) -> None:
    r = benchmark(lambda: extract(payload, mime=DOCX_MIME))
    assert len(r.content.text) > 0
